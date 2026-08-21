"""APA-formatted report generation — PDF, Word, CSV, Excel."""
import io, json
from datetime import datetime
from fastapi import APIRouter, Header, HTTPException
from fastapi.responses import StreamingResponse
from pydantic import BaseModel
from typing import Optional
from session.ephemeral_store import get_session, export_config

router = APIRouter()


class ReportRequest(BaseModel):
    format: str  # pdf | docx | csv | xlsx | json
    title: Optional[str] = "Statistical Analysis Report"
    include_audit: Optional[bool] = True
    include_hypotheses: Optional[bool] = True


@router.post("/generate")
def generate_report(req: ReportRequest,
                     session_token: str = Header(...), mode: str = Header(...)):
    session = get_session(session_token)
    if not session:
        raise HTTPException(404, "Session not found.")

    results = session.get("results", [])
    hypotheses = session.get("hypotheses", [])
    audit = session.get("audit", []) if req.include_audit else []
    alpha = session.get("alpha", 0.05)

    if req.format == "json":
        config = export_config(session_token)
        buf = io.BytesIO(json.dumps(config, indent=2).encode())
        return StreamingResponse(buf, media_type="application/json",
                                  headers={"Content-Disposition": "attachment; filename=analysis_config.json"})

    if req.format == "csv":
        import csv
        buf = io.StringIO()
        w = csv.writer(buf)
        w.writerow(["Test", "APA String", "p-value", "Reject H0", "Conclusion"])
        for r in results:
            if isinstance(r, dict) and "apa_string" in r:
                w.writerow([
                    r.get("test", ""),
                    r.get("apa_string", ""),
                    r.get("p_value", ""),
                    r.get("reject_h0", ""),
                    r.get("frequentist_conclusion", ""),
                ])
        out = io.BytesIO(buf.getvalue().encode())
        return StreamingResponse(out, media_type="text/csv",
                                  headers={"Content-Disposition": "attachment; filename=results.csv"})

    if req.format == "docx":
        return _generate_docx(req.title, results, hypotheses, audit, alpha, session)

    if req.format == "pdf":
        return _generate_pdf(req.title, results, hypotheses, audit, alpha, session)

    raise HTTPException(400, f"Unsupported format: {req.format}")


def _generate_docx(title, results, hypotheses, audit, alpha, session):
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()
    doc.add_heading(title, 0)

    # Metadata block
    meta = doc.add_paragraph()
    meta.add_run(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M')}\n")
    meta.add_run(f"α = {alpha}  |  Confidence Level: {session.get('confidence_level', 95)}%\n")
    meta.add_run(f"Total analyses: {len(results)}  |  Hypotheses evaluated: {len(hypotheses)}")

    # Hypotheses section
    if hypotheses:
        doc.add_heading("Hypothesis Evaluation", level=1)
        for h in hypotheses:
            doc.add_heading(f"H₁: {h.get('h1', '')}", level=2)
            p = doc.add_paragraph()
            p.add_run("H₀: ").bold = True
            p.add_run(h.get("h0", ""))
            if h.get("status"):
                verdict = doc.add_paragraph()
                verdict.add_run(f"Verdict: {h.get('verdict_label', '')} — ").bold = True
                verdict.add_run(h.get("formal_conclusion", ""))
            if h.get("apa_string"):
                doc.add_paragraph(f"Result: {h['apa_string']}")
            if h.get("plain_language"):
                doc.add_paragraph(h["plain_language"])

    # Statistical results
    doc.add_heading("Statistical Results", level=1)
    for r in results:
        if not isinstance(r, dict) or "test" not in r:
            continue
        doc.add_heading(r.get("test", "Analysis"), level=2)
        if "apa_string" in r:
            p = doc.add_paragraph()
            p.add_run("APA: ").bold = True
            p.add_run(r["apa_string"])
        if "frequentist_conclusion" in r:
            p = doc.add_paragraph()
            p.add_run("Conclusion: ").bold = True
            p.add_run(r["frequentist_conclusion"])
        if "plain_language" in r:
            doc.add_paragraph(r["plain_language"])
        if "caution" in r:
            p = doc.add_paragraph()
            run = p.add_run(f"⚠ {r['caution']}")
            run.font.color.rgb = RGBColor(0xB4, 0x50, 0x09)

    # Audit trail
    if audit:
        doc.add_heading("Audit Trail", level=1)
        doc.add_paragraph(
            "The following actions were recorded during this analysis session. "
            "This trail supports reproducibility and research transparency."
        )
        for entry in audit:
            p = doc.add_paragraph(style="List Bullet")
            p.add_run(f"[{entry['timestamp'][:19]}] ").bold = True
            p.add_run(f"{entry['action']}: {json.dumps(entry['details'])}")

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return StreamingResponse(buf,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f"attachment; filename=report.docx"})


def _generate_pdf(title, results, hypotheses, audit, alpha, session):
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import cm
    from reportlab.lib import colors
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, HRFlowable

    buf = io.BytesIO()
    doc = SimpleDocTemplate(buf, pagesize=A4,
                             leftMargin=2.5*cm, rightMargin=2.5*cm,
                             topMargin=2.5*cm, bottomMargin=2.5*cm)
    styles = getSampleStyleSheet()
    story = []

    story.append(Paragraph(title, styles["Title"]))
    story.append(Spacer(1, 0.3*cm))
    story.append(Paragraph(
        f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M')} | α = {alpha} | "
        f"Confidence: {session.get('confidence_level', 95)}%",
        styles["Normal"]
    ))
    story.append(HRFlowable(width="100%", thickness=1, color=colors.grey))
    story.append(Spacer(1, 0.5*cm))

    if hypotheses:
        story.append(Paragraph("Hypothesis Evaluation", styles["Heading1"]))
        for h in hypotheses:
            story.append(Paragraph(f"H₁: {h.get('h1','')}", styles["Heading2"]))
            story.append(Paragraph(f"H₀: {h.get('h0','')}", styles["Normal"]))
            if h.get("verdict_label"):
                story.append(Paragraph(
                    f"<b>{h.get('verdict_label','')}:</b> {h.get('formal_conclusion','')}",
                    styles["Normal"]
                ))
            if h.get("apa_string"):
                story.append(Paragraph(f"<i>{h['apa_string']}</i>", styles["Normal"]))
            story.append(Spacer(1, 0.3*cm))

    story.append(Paragraph("Statistical Results", styles["Heading1"]))
    for r in results:
        if not isinstance(r, dict) or "test" not in r:
            continue
        story.append(Paragraph(r.get("test", ""), styles["Heading2"]))
        if "apa_string" in r:
            story.append(Paragraph(f"<b>APA:</b> <i>{r['apa_string']}</i>", styles["Normal"]))
        if "frequentist_conclusion" in r:
            story.append(Paragraph(f"<b>Conclusion:</b> {r['frequentist_conclusion']}", styles["Normal"]))
        if "caution" in r:
            story.append(Paragraph(f"⚠ {r['caution']}", styles["Normal"]))
        story.append(Spacer(1, 0.3*cm))

    doc.build(story)
    buf.seek(0)
    return StreamingResponse(buf, media_type="application/pdf",
                              headers={"Content-Disposition": "attachment; filename=report.pdf"})


@router.get("/audit")
def export_audit(session_token: str = Header(...), mode: str = Header(...)):
    session = get_session(session_token)
    if not session:
        raise HTTPException(404, "Session not found.")
    return {"audit_trail": session.get("audit", []), "n_entries": len(session.get("audit", []))}


@router.get("/config")
def export_reproducibility_config(session_token: str = Header(...), mode: str = Header(...)):
    config = export_config(session_token)
    return config
